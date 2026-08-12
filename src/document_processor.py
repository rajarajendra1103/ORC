import os
import io
from pathlib import Path
from PIL import Image
import numpy as np

# Optional file format handlers
try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import fitz # PyMuPDF
except ImportError:
    fitz = None

try:
    import docx
except ImportError:
    docx = None

try:
    import openpyxl
    import pandas as pd
except ImportError:
    openpyxl = None
    pd = None

from src.models.trocr_extractor import TrOCRExtractor
from src.models.layoutlm_analyzer import LayoutLMv3Analyzer
from src.models.donut_parser import DonutParser

class DocumentProcessor:
    """
    Unified Multi-Format Document Processor
    Handles Images, PDFs, DOCX, and XLSX files.
    Runs TrOCR, LayoutLMv3, and Donut Transformer pipeline to produce Unified Document Knowledge.
    """
    def __init__(self):
        self.trocr = TrOCRExtractor()
        self.layoutlm = LayoutLMv3Analyzer()
        self.donut = DonutParser()

    def process_file(self, file_path_or_bytes, filename=None):
        """
        Process any supported document format and return Unified Document Knowledge.
        """
        if filename is None:
            if isinstance(file_path_or_bytes, (str, Path)):
                filename = str(file_path_or_bytes)
            else:
                filename = "document.png"

        ext = Path(filename).suffix.lower()
        
        pages_knowledge = []

        if ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".webp"]:
            # Single Image Processing
            page_data = self._process_image_input(file_path_or_bytes, page_num=1)
            pages_knowledge.append(page_data)

        elif ext == ".pdf":
            # PDF Processing (Native Text + Image Rendering)
            pages_knowledge = self._process_pdf(file_path_or_bytes)

        elif ext in [".docx", ".doc"]:
            # Word Document Processing
            pages_knowledge = self._process_docx(file_path_or_bytes)

        elif ext in [".xlsx", ".xls"]:
            # Excel Spreadsheet Processing
            pages_knowledge = self._process_xlsx(file_path_or_bytes)

        else:
            # Fallback to image loading or plain text reading
            page_data = self._process_image_input(file_path_or_bytes, page_num=1)
            pages_knowledge.append(page_data)

        # Aggregate across pages into Unified Document Knowledge
        all_text = "\n\n--- PAGE BREAK ---\n\n".join([p["text"] for p in pages_knowledge if p["text"].strip()])
        all_blocks = []
        for p in pages_knowledge:
            all_blocks.extend(p.get("layout_blocks", []))

        unified_structured = self.donut.parse_structured_data(all_text, all_blocks)["structured_json"]

        return {
            "filename": os.path.basename(filename),
            "file_type": ext,
            "total_pages": len(pages_knowledge),
            "unified_text": all_text,
            "structured_knowledge": unified_structured,
            "pages": pages_knowledge
        }

    def _process_image_input(self, image_input, page_num=1):
        if isinstance(image_input, (str, Path)):
            pil_img = Image.open(image_input).convert("RGB")
        elif isinstance(image_input, bytes):
            pil_img = Image.open(io.BytesIO(image_input)).convert("RGB")
        elif isinstance(image_input, Image.Image):
            pil_img = image_input.convert("RGB")
        else:
            pil_img = Image.fromarray(image_input).convert("RGB")

        w, h = pil_img.size

        # 1. TrOCR Text Extraction
        trocr_res = self.trocr.extract_text_and_boxes(pil_img)
        
        # 2. LayoutLMv3 Layout Analysis
        layout_res = self.layoutlm.analyze_layout(trocr_res["lines"], image_width=w, image_height=h)

        # 3. Donut Transformer Structured Understanding
        donut_res = self.donut.parse_structured_data(trocr_res["full_text"], layout_res["blocks"])

        return {
            "page_number": page_num,
            "image": pil_img,
            "width": w,
            "height": h,
            "text": trocr_res["full_text"],
            "ocr_lines": trocr_res["lines"],
            "layout_blocks": layout_res["blocks"],
            "layout_summary": layout_res["layout_summary"],
            "structured_json": donut_res["structured_json"]
        }

    def _process_pdf(self, pdf_input):
        pages = []
        pdf_bytes = pdf_input if isinstance(pdf_input, bytes) else open(pdf_input, "rb").read()

        # Try PyMuPDF (fitz) for page rendering & text extraction
        if fitz:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            for page_idx in range(len(doc)):
                page = doc[page_idx]
                pix = page.get_pixmap(dpi=150)
                pil_img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
                
                # Check native text
                native_text = page.get_text().strip()
                if len(native_text) > 30:
                    w, h = pil_img.size
                    layout_res = self.layoutlm.analyze_layout([], image_width=w, image_height=h)
                    donut_res = self.donut.parse_structured_data(native_text)
                    pages.append({
                        "page_number": page_idx + 1,
                        "image": pil_img,
                        "width": w,
                        "height": h,
                        "text": native_text,
                        "ocr_lines": [],
                        "layout_blocks": [],
                        "layout_summary": layout_res["layout_summary"],
                        "structured_json": donut_res["structured_json"]
                    })
                else:
                    # Run OCR vision pipeline on rendered page
                    pages.append(self._process_image_input(pil_img, page_num=page_idx + 1))
            return pages

        # Fallback to PyPDF text extraction
        if pypdf:
            reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
            for page_idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                dummy_img = Image.new("RGB", (800, 1000), color=(255, 255, 255))
                donut_res = self.donut.parse_structured_data(text)
                pages.append({
                    "page_number": page_idx + 1,
                    "image": dummy_img,
                    "width": 800,
                    "height": 1000,
                    "text": text,
                    "ocr_lines": [],
                    "layout_blocks": [],
                    "layout_summary": {"header_count": 0, "paragraph_count": 1, "table_region_count": 0},
                    "structured_json": donut_res["structured_json"]
                })
            return pages

        return pages

    def _process_docx(self, docx_input):
        if not docx:
            raise ImportError("python-docx is required for DOCX processing.")

        docx_bytes = docx_input if isinstance(docx_input, bytes) else open(docx_input, "rb").read()
        doc = docx.Document(io.BytesIO(docx_bytes))
        
        full_text_lines = []
        for p in doc.paragraphs:
            if p.text.strip():
                full_text_lines.append(p.text.strip())

        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_cells))
            full_text_lines.append("\n".join(table_text))

        full_text = "\n".join(full_text_lines)
        dummy_img = Image.new("RGB", (800, 1000), color=(248, 249, 250))
        donut_res = self.donut.parse_structured_data(full_text)

        return [{
            "page_number": 1,
            "image": dummy_img,
            "width": 800,
            "height": 1000,
            "text": full_text,
            "ocr_lines": [],
            "layout_blocks": [],
            "layout_summary": {"header_count": 0, "paragraph_count": len(full_text_lines), "table_region_count": len(doc.tables)},
            "structured_json": donut_res["structured_json"]
        }]

    def _process_xlsx(self, xlsx_input):
        if pd is None:
            raise ImportError("openpyxl and pandas are required for XLSX processing.")

        xlsx_bytes = xlsx_input if isinstance(xlsx_input, bytes) else open(xlsx_input, "rb").read()
        excel_file = pd.ExcelFile(io.BytesIO(xlsx_bytes))
        
        full_text_lines = []
        for sheet_name in excel_file.sheet_names:
            df = excel_file.parse(sheet_name)
            full_text_lines.append(f"=== Sheet: {sheet_name} ===")
            full_text_lines.append(df.to_string())

        full_text = "\n\n".join(full_text_lines)
        dummy_img = Image.new("RGB", (800, 1000), color=(240, 245, 250))
        donut_res = self.donut.parse_structured_data(full_text)

        return [{
            "page_number": 1,
            "image": dummy_img,
            "width": 800,
            "height": 1000,
            "text": full_text,
            "ocr_lines": [],
            "layout_blocks": [],
            "layout_summary": {"header_count": len(excel_file.sheet_names), "paragraph_count": 0, "table_region_count": len(excel_file.sheet_names)},
            "structured_json": donut_res["structured_json"]
        }]
